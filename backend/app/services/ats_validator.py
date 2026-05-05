"""
ATS Validator Service for checking resume compliance with ATS systems.
"""
import re
import os
import fitz  # PyMuPDF
import docx
from typing import Dict, List, Tuple, Any, Optional
import logging

logger = logging.getLogger(__name__)

class ATSValidator:
    """
    Service to validate if a resume is ATS compliant based on various rules.
    """
    
    # Standard fonts that are ATS-friendly
    STANDARD_FONTS = [
        'arial', 'helvetica', 'calibri', 'times new roman', 'times', 
        'georgia', 'verdana', 'trebuchet', 'tahoma', 'garamond'
    ]
    
    # Common section headers in resumes
    RECOGNIZED_HEADERS = [
        'summary', 'profile', 'objective', 'experience', 'work experience', 
        'employment', 'education', 'skills', 'technical skills', 'certifications',
        'projects', 'achievements', 'awards', 'publications', 'languages',
        'interests', 'references', 'contact', 'personal information'
    ]
    
    # Date format regex patterns
    DATE_PATTERNS = [
        r'\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]* \d{4}\b',  # Month Year
        r'\b\d{1,2}/\d{1,2}/\d{2,4}\b',  # MM/DD/YYYY
        r'\b\d{1,2}-\d{1,2}-\d{2,4}\b',  # MM-DD-YYYY
        r'\b\d{4}\b',  # Just year
        r'\b(present|current|now)\b',  # Present indicators
        r'\b\d{4}-\d{4}\b',  # Year range
        r'\b\d{4} - \d{4}\b',  # Year range with spaces
        r'\b\d{4} to \d{4}\b',  # Year range with "to"
        r'\b\d{4}-present\b',  # Year to present
        r'\b\d{4} - present\b',  # Year to present with spaces
    ]
    
    def __init__(self):
        """Initialize the ATS validator."""
        pass
    
    def validate_resume(self, resume_input: str, job_description: Optional[str] = None) -> Dict[str, Any]:
        """
        Validate a resume file or resume text against ATS compliance rules.
        
        Args:
            resume_input: Path to the resume file (PDF or DOCX) or resume text
            job_description: Optional job description when validating text input
            
        Returns:
            Dictionary with validation results for each rule and overall score
        """
        file_ext = os.path.splitext(resume_input)[1].lower()

        try:
            if file_ext == '.pdf':
                return self._validate_pdf(resume_input)
            elif file_ext == '.docx':
                return self._validate_docx(resume_input)
            elif file_ext:
                return {
                    "error": f"Unsupported file format: {file_ext}. Please upload a PDF or DOCX file.",
                    "overall_score": 0,
                    "passed": False
                }
            return self._validate_text(resume_input, job_description)
        except Exception as e:
            logger.error(f"Error validating resume: {str(e)}")
            return {
                "error": f"Error processing file: {str(e)}",
                "overall_score": 0,
                "passed": False
            }

    def _validate_text(self, resume_text: str, job_description: Optional[str] = None) -> Dict[str, Any]:
        """Validate resume text against basic ATS rules."""
        results = self._validate_content(resume_text)
        overall_score = self._calculate_score(results)
        suggestions = []
        if job_description:
            suggestions.append("Include keywords from the job description where relevant.")

        return {
            "overall_score": overall_score,
            "recognized_headers": results.get("recognized_section_headers", False),
            "proper_date_formats": results.get("proper_date_formats", False),
            "content_score": overall_score,
            "suggestions": suggestions
        }
    
    def _validate_pdf(self, file_path: str) -> Dict[str, Any]:
        """Validate a PDF resume file."""
        try:
            # Open the PDF
            doc = fitz.open(file_path)
            
            # Extract text from all pages
            full_text = ""
            for page in doc:
                full_text += page.get_text()
            
            # Check if text is extractable
            is_text_extractable = len(full_text.strip()) > 0
            
            # Get validation results
            results = self._validate_content(full_text)
            
            # PDF-specific checks
            results["single_column"] = self._check_single_column_pdf(doc)
            results["no_images_tables_textboxes"] = self._check_no_images_tables_pdf(doc)
            results["text_extractable"] = is_text_extractable
            
            # Calculate overall score
            overall_score = self._calculate_score(results)
            
            # Add overall results
            results["overall_score"] = overall_score
            results["passed"] = overall_score >= 80
            
            return results
            
        except Exception as e:
            logger.error(f"Error validating PDF: {str(e)}")
            raise
    
    def _validate_docx(self, file_path: str) -> Dict[str, Any]:
        """Validate a DOCX resume file."""
        try:
            # Open the DOCX
            doc = docx.Document(file_path)
            
            # Extract text
            full_text = "\n".join([para.text for para in doc.paragraphs])
            
            # Check if text is extractable
            is_text_extractable = len(full_text.strip()) > 0
            
            # Get validation results
            results = self._validate_content(full_text)
            
            # DOCX-specific checks
            results["single_column"] = self._check_single_column_docx(doc)
            results["no_images_tables_textboxes"] = self._check_no_images_tables_docx(doc)
            results["text_extractable"] = is_text_extractable
            
            # Check fonts
            results["standard_fonts"] = self._check_standard_fonts_docx(doc)
            
            # Calculate overall score
            overall_score = self._calculate_score(results)
            
            # Add overall results
            results["overall_score"] = overall_score
            results["passed"] = overall_score >= 80
            
            return results
            
        except Exception as e:
            logger.error(f"Error validating DOCX: {str(e)}")
            raise
    
    def _validate_content(self, text: str) -> Dict[str, bool]:
        """Validate the content of the resume text."""
        results = {}
        
        # Check for recognized section headers
        results["recognized_section_headers"] = self._check_recognized_headers(text)
        
        # Check date formats
        results["proper_date_formats"] = self._check_date_formats(text)
        
        # For PDF we'll set this as default and override in specific methods
        results["standard_fonts"] = True
        
        return results
    
    def _check_single_column_pdf(self, doc: fitz.Document) -> bool:
        """Check if the PDF uses a single column layout."""
        # This is a simplified heuristic - checking if text blocks are roughly aligned
        for page in doc:
            blocks = page.get_text("blocks")
            if len(blocks) < 2:
                continue
                
            # Get x-coordinates of blocks
            x_coords = [block[0] for block in blocks]
            
            # If there are multiple distinct x-coordinate clusters, likely multi-column
            x_clusters = self._cluster_values(x_coords, threshold=50)
            if len(x_clusters) > 1:
                return False
                
        return True
    
    def _check_single_column_docx(self, doc: docx.Document) -> bool:
        """Check if the DOCX uses a single column layout."""
        # DOCX files are typically single column unless sections are used
        # This is a simplified check
        for section in doc.sections:
            if section.start_type != 0:  # Not continuous section
                return False
        return True
    
    def _check_no_images_tables_pdf(self, doc: fitz.Document) -> bool:
        """Check if the PDF has no images or tables."""
        for page in doc:
            # Check for images
            images = page.get_images(full=True)
            if len(images) > 0:
                return False
                
            # Simplified table detection - look for grid-like structures
            # This is not perfect but gives a reasonable approximation
            lines = page.get_drawings()
            horizontal_lines = 0
            vertical_lines = 0
            
            for line in lines:
                if "rect" in line:
                    rect = line["rect"]
                    if abs(rect[1] - rect[3]) < 2:  # Horizontal line
                        horizontal_lines += 1
                    if abs(rect[0] - rect[2]) < 2:  # Vertical line
                        vertical_lines += 1
            
            # If we have multiple horizontal and vertical lines, likely a table
            if horizontal_lines > 3 and vertical_lines > 3:
                return False
                
        return True
    
    def _check_no_images_tables_docx(self, doc: docx.Document) -> bool:
        """Check if the DOCX has no images or tables."""
        # Check for tables
        if len(doc.tables) > 0:
            return False
            
        # Check for images (simplified)
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                return False
                
        # Check for text boxes (simplified)
        if hasattr(doc, 'inline_shapes') and len(doc.inline_shapes) > 0:
            return False
            
        return True
    
    def _check_recognized_headers(self, text: str) -> bool:
        """Check if the resume has recognized section headers."""
        text_lower = text.lower()
        
        # Check for presence of common section headers
        header_count = 0
        for header in self.RECOGNIZED_HEADERS:
            # Look for headers that are on their own line or followed by a colon
            pattern = fr'\b{re.escape(header)}\b\s*:|\n\s*\b{re.escape(header)}\b\s*\n'
            if re.search(pattern, text_lower, re.IGNORECASE):
                header_count += 1
                
        # If we found at least 3 recognized headers, consider it compliant
        return header_count >= 3
    
    def _check_date_formats(self, text: str) -> bool:
        """Check if the resume uses proper date formats."""
        # Look for date patterns
        date_matches = 0
        for pattern in self.DATE_PATTERNS:
            matches = re.findall(pattern, text.lower())
            date_matches += len(matches)
            
        # If we found at least 2 properly formatted dates, consider it compliant
        return date_matches >= 2
    
    def _check_standard_fonts_docx(self, doc: docx.Document) -> bool:
        """Check if the DOCX uses standard fonts."""
        non_standard_fonts = 0
        total_runs = 0
        
        for para in doc.paragraphs:
            for run in para.runs:
                total_runs += 1
                if run.font.name and run.font.name.lower() not in self.STANDARD_FONTS:
                    non_standard_fonts += 1
                    
        # If more than 20% of text runs use non-standard fonts, fail this check
        if total_runs == 0:
            return True
        return (non_standard_fonts / total_runs) < 0.2
    
    def _calculate_score(self, results: Dict[str, bool]) -> int:
        """Calculate the overall compliance score."""
        # Remove non-rule keys
        rule_results = {k: v for k, v in results.items() 
                       if k not in ['overall_score', 'passed', 'error']}
        
        # Count passed rules
        passed_rules = sum(1 for v in rule_results.values() if v)
        total_rules = len(rule_results)
        
        # Calculate percentage
        if total_rules == 0:
            return 0
        return int((passed_rules / total_rules) * 100)
    
    def _cluster_values(self, values: List[float], threshold: float) -> List[List[float]]:
        """Group values into clusters based on proximity."""
        if not values:
            return []
            
        # Sort values
        sorted_values = sorted(values)
        
        # Initialize clusters
        clusters = [[sorted_values[0]]]
        
        # Group values
        for value in sorted_values[1:]:
            if value - clusters[-1][-1] <= threshold:
                clusters[-1].append(value)
            else:
                clusters.append([value])
                
        return clusters
