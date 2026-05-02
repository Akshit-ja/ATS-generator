import io
import os
import uuid
from typing import Dict, Any
from pathlib import Path
import tempfile
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch


class DocumentGeneratorService:
    """Service for generating resume documents in various formats"""
    
    def __init__(self):
        self.temp_dir = Path(tempfile.gettempdir()) / "resume_generator"
        self.temp_dir.mkdir(exist_ok=True)
    
    def generate_docx(self, resume_data: Dict[str, Any]) -> str:
        """Generate a DOCX resume and return the file path"""
        try:
            doc = Document()
            
            # Title/Name
            title = doc.add_heading('Resume', 0)
            title.alignment = 1  # Center alignment
            
            # Professional Summary
            if resume_data.get('professional_summary'):
                doc.add_heading('Professional Summary', level=1)
                doc.add_paragraph(resume_data['professional_summary'])
                doc.add_paragraph('')
            
            # Technical Skills
            if resume_data.get('technical_skills'):
                doc.add_heading('Technical Skills', level=1)
                doc.add_paragraph(resume_data['technical_skills'])
                doc.add_paragraph('')
            
            # Work Experience
            if resume_data.get('work_experience'):
                doc.add_heading('Work Experience', level=1)
                doc.add_paragraph(resume_data['work_experience'])
                doc.add_paragraph('')
            
            # Education
            if resume_data.get('education'):
                doc.add_heading('Education', level=1)
                doc.add_paragraph(resume_data['education'])
                doc.add_paragraph('')
            
            # Additional Skills
            if resume_data.get('additional_skills'):
                doc.add_heading('Additional Skills', level=1)
                doc.add_paragraph(resume_data['additional_skills'])
            
            # Save document
            filename = f"resume_{uuid.uuid4()}.docx"
            file_path = self.temp_dir / filename
            doc.save(str(file_path))
            
            return str(file_path)
            
        except Exception as e:
            raise Exception(f"Error generating DOCX: {str(e)}")
    
    def generate_pdf(self, resume_data: Dict[str, Any]) -> str:
        """Generate a PDF resume and return the file path"""
        try:
            filename = f"resume_{uuid.uuid4()}.pdf"
            file_path = self.temp_dir / filename
            
            # Create PDF document
            doc = SimpleDocTemplate(str(file_path), pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Title style
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=1  # Center alignment
            )
            
            # Add title
            story.append(Paragraph("Resume", title_style))
            story.append(Spacer(1, 20))
            
            # Professional Summary
            if resume_data.get('professional_summary'):
                story.append(Paragraph("Professional Summary", styles['Heading2']))
                story.append(Paragraph(resume_data['professional_summary'], styles['Normal']))
                story.append(Spacer(1, 12))
            
            # Technical Skills
            if resume_data.get('technical_skills'):
                story.append(Paragraph("Technical Skills", styles['Heading2']))
                story.append(Paragraph(resume_data['technical_skills'], styles['Normal']))
                story.append(Spacer(1, 12))
            
            # Work Experience
            if resume_data.get('work_experience'):
                story.append(Paragraph("Work Experience", styles['Heading2']))
                story.append(Paragraph(resume_data['work_experience'], styles['Normal']))
                story.append(Spacer(1, 12))
            
            # Education
            if resume_data.get('education'):
                story.append(Paragraph("Education", styles['Heading2']))
                story.append(Paragraph(resume_data['education'], styles['Normal']))
                story.append(Spacer(1, 12))
            
            # Additional Skills
            if resume_data.get('additional_skills'):
                story.append(Paragraph("Additional Skills", styles['Heading2']))
                story.append(Paragraph(resume_data['additional_skills'], styles['Normal']))
            
            # Build PDF
            doc.build(story)
            
            return str(file_path)
            
        except Exception as e:
            raise Exception(f"Error generating PDF: {str(e)}")
    
    def cleanup_file(self, file_path: str):
        """Clean up temporary file"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception as e:
            print(f"Warning: Could not clean up file {file_path}: {str(e)}")